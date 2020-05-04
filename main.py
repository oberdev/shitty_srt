from services.scheduling import SchedullingService
from utils.task import PeriodicTask, AperiodicTask
from flask import Flask, Markup
from flask import render_template, jsonify
import json
from input import PERIODIC_TASKS, APERIODIC_TASKS, HYPERPERIODS_COUNT
from copy import deepcopy
from docx import Document
from typing import List
import numpy as np

INPUT_FILE_NAME = "tasks.json"


def find_by_id(id):
    def comparator(task):
        return task.id == id

    return comparator


def toFixed(numObj, digits=0):
    return f"{numObj:.{digits}f}"


def read_tasks_from_file(input_file_name: str):
    input_tasks = {}
    periodic_tasks = []
    aperiodic_tasks = []
    with open(input_file_name) as file:
        input_tasks = json.loads(file.read())
    for task in input_tasks["periodic"]:
        periodic_tasks.append(PeriodicTask(task["id"], task["period"], task["exec_time"]))
    for task in input_tasks["aperiodic"]:
        aperiodic_tasks.append(AperiodicTask(task["id"], task["period"], task["exec_time"]))
    return periodic_tasks, aperiodic_tasks


def create_cpu_trace(document_name: str, tasks: List[PeriodicTask], result: dict):
    document = Document()
    tasks_ids = [task.id for task in tasks]
    table = document.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Период'
    hdr_cells[2].text = 'λ'
    hdr_cells[3].text = 'Время выполнения'
    hdr_cells[4].text = 'Среднее время отклика'
    hdr_cells[5].text = 'Макс. время отклика'
    hdr_cells[6].text = 'D'
    for key, task in enumerate(tasks):
        row_cells = table.add_row().cells
        row_cells[0].text = str(task.id)
        row_cells[1].text = toFixed(task.deadline_param / 1000, 4) if isinstance(task, AperiodicTask) else str(
            task.deadline / 1000)
        row_cells[2].text = toFixed(1 / (task.deadline_param / 1000), 4) if isinstance(task, AperiodicTask) else ''
        row_cells[3].text = str(task.exec_time / 1000)
        row_cells[4].text = toFixed(result[task.id]["mean"] / 1000, 1)
        row_cells[5].text = toFixed(result[task.id]["max"] / 1000, 1)
        row_cells[6].text = row_cells[1].text
    document.save(f'{document_name}.docx')


def create_tables_trace(document_name, trace, tasks):
    document = Document()
    tasks_ids = [task.id for task in tasks]
    for key in tasks_ids:
        task: PeriodicTask = list(filter(find_by_id(key), tasks))[0]
        document.add_heading(
            f'{"Апериодическая" if isinstance(task, AperiodicTask) else "Периодическая"} задача №{key} Период: {int(task.deadline_param) / 1000 if isinstance(task, AperiodicTask) else (task.deadline) / 1000} Время исполнения: {(task.exec_time) / 1000}',
            level=2)
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Номер вызова'
        hdr_cells[1].text = 'Время вызова'
        hdr_cells[2].text = 'Время запуска'
        hdr_cells[3].text = 'Время отклика'
        for task_num, task_trace in enumerate(trace[key]):
            row_cells = table.add_row().cells
            row_cells[0].text = str(task_num + 1)
            row_cells[1].text = str(task_trace['came_moment'] / 1000)
            row_cells[2].text = str(task_trace['execute_moment'] / 1000)
            row_cells[3].text = str(task_trace['response_time'] / 1000)
        document.add_page_break()

    document.save(f'{document_name}.docx')


# periodic_tasks, aperiodic_tasks = read_tasks_from_file(INPUT_FILE_NAME)
# tasks = periodic_tasks + aperiodic_tasks

app = Flask(__name__)


@app.route("/edf")
def edf():
    title, trace_data, responses, results, error = SchedullingService(periodic_tasks, deepcopy(aperiodic_tasks),
                                                                      HYPERPERIODS_COUNT).run('edf')
    create_cpu_trace('edf_cpu', tasks, results)
    create_tables_trace('edf', responses)
    with open('edf_stat.json', 'w') as f:
        json.dump(results, f)
    with open('edf_out.json', 'w') as f:
        json.dump(trace_data, f)
    return render_template("index.html", title=title, trace_data=Markup(trace_data))


@app.route("/rm")
def rm():
    seed = 0
    # seed = 6 task 1
    # seed = 103 task 2
    seed = 0
    while True:
        np.random.seed(seed)
        try:
            periodic_tasks, aperiodic_tasks = read_tasks_from_file(INPUT_FILE_NAME)
            tasks = periodic_tasks + aperiodic_tasks
            title, trace_data, responses, results, error = SchedullingService(periodic_tasks, deepcopy(aperiodic_tasks),
                                                                              HYPERPERIODS_COUNT).run('rm')
            create_cpu_trace('rm_cpu', tasks, results)
            create_tables_trace('rm', responses, tasks)
            with open('rm_stat.json', 'w') as f:
                json.dump(results, f)
            with open('rm_out.json', 'w') as f:
                json.dump(trace_data, f)
            print(seed)
            return render_template("index.html", title=title, trace_data=Markup(trace_data))
        except KeyError:
            seed += 1
            continue


if __name__ == "__main__":
    app.run()
